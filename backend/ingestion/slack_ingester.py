import json
import logging
from datetime import datetime, timezone
from pathlib import Path
from typing import Any

from sqlalchemy import select
from sqlalchemy.ext.asyncio import AsyncSession

from backend.knowledge.models import Document
from backend.processing.embedder import DocumentEmbedder

logger = logging.getLogger(__name__)


class SlackExportIngester:
    """Ingests a Slack export directory into Cortex documents with embeddings."""

    def __init__(self, export_dir: str | Path):
        self.export_dir = Path(export_dir)
        self.users: dict[str, dict] = {}        # user_id -> user data
        self.channels: dict[str, dict] = {}      # channel_id -> channel data
        self.channel_names: dict[str, str] = {}  # channel directory name -> display name
        self.embedder = DocumentEmbedder()
        self.stats = {
            "messages_processed": 0,
            "threads_processed": 0,
            "attachments_processed": 0,
            "documents_created": 0,
            "documents_skipped_existing": 0,
            "errors": 0,
        }

    def _load_users(self) -> None:
        """Load users.json and build lookup."""
        users_file = self.export_dir / "users.json"
        if not users_file.exists():
            logger.warning("users.json not found in export directory")
            return
        with open(users_file) as f:
            users_list = json.load(f)
        for user in users_list:
            uid = user.get("id", "")
            self.users[uid] = {
                "name": user.get("real_name") or user.get("name", "unknown"),
                "role": user.get("profile", {}).get("title", ""),
            }

    def _load_channels(self) -> None:
        """Load channels.json and build lookup."""
        channels_file = self.export_dir / "channels.json"
        if not channels_file.exists():
            logger.warning("channels.json not found in export directory")
            return
        with open(channels_file) as f:
            channels_list = json.load(f)
        for ch in channels_list:
            cid = ch.get("id", "")
            name = ch.get("name", "")
            self.channels[cid] = ch
            self.channel_names[name] = name

    def _get_user_name(self, user_id: str) -> str:
        return self.users.get(user_id, {}).get("name", user_id)

    def _get_user_role(self, user_id: str) -> str:
        return self.users.get(user_id, {}).get("role", "")

    def _build_slack_link(self, channel_name: str, ts: str) -> str:
        """Build a Slack deep link from channel name and timestamp."""
        ts_formatted = ts.replace(".", "")
        return f"slack://channel?team=T0&id={channel_name}&message={ts_formatted}"

    def _extract_attachment_text(self, file_info: dict) -> str | None:
        """Extract text from file attachments (PDF, DOCX, CSV, images)."""
        file_path_str = file_info.get("local_path") or file_info.get("url_private_download")
        if not file_path_str:
            return None

        # Check if the file exists locally in the export
        file_path = self.export_dir / "__uploads" / file_info.get("id", "")
        if not file_path.exists():
            # Try constructing from the file info
            name = file_info.get("name", "")
            file_path = self.export_dir / "__uploads" / name
            if not file_path.exists():
                logger.debug("Attachment file not found locally: %s", name)
                return None

        mimetype = file_info.get("mimetype", "")
        name = file_info.get("name", "").lower()

        try:
            if mimetype == "application/pdf" or name.endswith(".pdf"):
                return self._extract_pdf(file_path)
            elif mimetype in (
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            ) or name.endswith(".docx"):
                return self._extract_docx(file_path)
            elif mimetype == "text/csv" or name.endswith(".csv"):
                return self._extract_csv(file_path)
            elif mimetype.startswith("image/"):
                return self._extract_image_text(file_path)
        except Exception as e:
            logger.warning("Failed to extract text from %s: %s", name, e)
            self.stats["errors"] += 1

        return None

    def _extract_pdf(self, path: Path) -> str:
        import pdfplumber
        text_parts = []
        with pdfplumber.open(path) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)
        return "\n".join(text_parts)

    def _extract_docx(self, path: Path) -> str:
        from docx import Document as DocxDocument
        doc = DocxDocument(str(path))
        return "\n".join(p.text for p in doc.paragraphs if p.text.strip())

    def _extract_csv(self, path: Path) -> str:
        import pandas as pd
        df = pd.read_csv(path, nrows=500)
        return df.to_string(index=False, max_rows=100)

    def _extract_image_text(self, path: Path) -> str:
        import pytesseract
        from PIL import Image
        img = Image.open(path)
        return pytesseract.image_to_string(img)

    def _parse_messages(self, channel_name: str, messages: list[dict]) -> list[dict]:
        """Parse messages, grouping threads together."""
        # Group messages by thread_ts
        threads: dict[str, list[dict]] = {}
        standalone: list[dict] = []

        for msg in messages:
            if msg.get("subtype") in (
                "channel_join", "channel_leave", "channel_topic", "channel_purpose",
                "bot_message", "file_comment", "tombstone", "thread_broadcast",
            ):
                continue
            text = (msg.get("text") or "").strip()
            if not text and not msg.get("files"):
                continue

            thread_ts = msg.get("thread_ts")
            ts = msg.get("ts", "")

            if thread_ts and thread_ts != ts:
                # This is a reply in a thread
                threads.setdefault(thread_ts, []).append(msg)
            elif thread_ts and thread_ts == ts:
                # This is the parent of a thread
                threads.setdefault(ts, []).insert(0, msg)
            else:
                standalone.append(msg)

        documents = []

        # Process standalone messages
        for msg in standalone:
            ts = msg.get("ts", "")
            user_id = msg.get("user", "") or msg.get("username", msg.get("bot_id", "bot"))
            text = msg.get("text", "")

            # Handle file attachments
            attachment_texts = []
            for file_info in msg.get("files", []):
                extracted = self._extract_attachment_text(file_info)
                if extracted:
                    attachment_texts.append(f"[Attachment: {file_info.get('name', 'file')}]\n{extracted}")
                    self.stats["attachments_processed"] += 1

            full_content = text
            if attachment_texts:
                full_content += "\n\n" + "\n\n".join(attachment_texts)

            try:
                # Naive UTC — asyncpg rejects tz-aware datetimes on naive columns
                created_at = (
                    datetime.fromtimestamp(float(ts), tz=timezone.utc)
                    .replace(tzinfo=None)
                    if ts else None
                )
            except (ValueError, OSError):
                created_at = None

            documents.append({
                "content": full_content,
                "source_type": "slack",
                "source_id": ts,
                "source_link": self._build_slack_link(channel_name, ts),
                "source_label": f"#{channel_name}",
                "channel_or_project": channel_name,
                "author_name": self._get_user_name(user_id),
                "author_role": self._get_user_role(user_id),
                "created_at": created_at,
            })
            self.stats["messages_processed"] += 1

        # Process threads — group all replies into one document
        for parent_ts, thread_msgs in threads.items():
            thread_parts = []
            first_user = ""
            for i, msg in enumerate(thread_msgs):
                user_id = msg.get("user", "")
                if i == 0:
                    first_user = user_id
                user_name = self._get_user_name(user_id)
                text = msg.get("text", "")
                thread_parts.append(f"[{user_name}]: {text}")

                for file_info in msg.get("files", []):
                    extracted = self._extract_attachment_text(file_info)
                    if extracted:
                        thread_parts.append(f"[Attachment: {file_info.get('name', 'file')}]\n{extracted}")
                        self.stats["attachments_processed"] += 1

            full_content = "\n".join(thread_parts)

            documents.append({
                "content": full_content,
                "source_type": "slack",
                "source_id": parent_ts,
                "source_link": self._build_slack_link(channel_name, parent_ts),
                "source_label": f"#{channel_name} (thread)",
                "channel_or_project": channel_name,
                "author_name": self._get_user_name(first_user),
                "author_role": self._get_user_role(first_user),
                "created_at": (
                    datetime.fromtimestamp(float(parent_ts), tz=timezone.utc)
                    .replace(tzinfo=None)
                    if parent_ts else None
                ),
            })
            self.stats["threads_processed"] += 1

        return documents

    async def ingest(self, db: AsyncSession) -> dict[str, Any]:
        """Run the full ingestion pipeline.

        1. Load users and channels metadata
        2. Parse all channel message files
        3. Create Document records in PostgreSQL
        4. Generate embeddings and store in ChromaDB

        Returns dict with ingestion stats.
        """
        logger.info("Starting Slack export ingestion from %s", self.export_dir)

        # Step 1: Load metadata
        self._load_users()
        self._load_channels()

        # Step 2: Parse all channels
        all_docs: list[dict] = []
        for channel_dir in sorted(self.export_dir.iterdir()):
            if not channel_dir.is_dir() or channel_dir.name.startswith(("_", ".")):
                continue

            channel_name = channel_dir.name
            logger.info("Processing channel: #%s", channel_name)

            for json_file in sorted(channel_dir.glob("*.json")):
                try:
                    with open(json_file) as f:
                        messages = json.load(f)
                    docs = self._parse_messages(channel_name, messages)
                    all_docs.extend(docs)
                except Exception as e:
                    logger.error("Error processing %s: %s", json_file, e)
                    self.stats["errors"] += 1

        logger.info("Parsed %d documents from Slack export", len(all_docs))

        # Step 3: Store in PostgreSQL — skipping messages already ingested
        # (same source_type + source_id), so re-uploading an export is
        # idempotent instead of duplicating every document.
        source_ids = {d["source_id"] for d in all_docs if d.get("source_id")}
        existing: set[str] = set()
        if source_ids:
            result = await db.execute(
                select(Document.source_id).where(
                    Document.source_type == "slack",
                    Document.source_id.in_(source_ids),
                )
            )
            existing = {row[0] for row in result.all()}

        created_docs: list[Document] = []
        for doc_data in all_docs:
            if doc_data.get("source_id") in existing:
                self.stats["documents_skipped_existing"] += 1
                continue
            doc = Document(**doc_data)
            db.add(doc)
            created_docs.append(doc)

        await db.flush()
        for doc in created_docs:
            await db.refresh(doc)

        self.stats["documents_created"] = len(created_docs)
        logger.info("Stored %d documents in PostgreSQL", len(created_docs))

        # Step 4: Generate embeddings and store in ChromaDB
        if created_docs:
            await self.embedder.embed_documents(
                db, document_ids=[d.id for d in created_docs]
            )

        logger.info("Slack ingestion complete: %s", self.stats)
        return self.stats
